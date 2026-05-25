"""DocCheck 测试辅助函数 — 构造测试文档"""

from io import BytesIO
from docx import Document


def make_docx(content: str, filename: str = "test.docx") -> BytesIO:
    """Create a .docx file in memory with given text content.

    Args:
        content: Plain text content. Lines starting with # are treated as headings.
        filename: Display filename (not important for BytesIO).

    Returns:
        BytesIO stream suitable for upload.
    """
    doc = Document()
    for line in content.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            # Heading level 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    buf.name = filename
    return buf


def make_compliant_report_docx() -> BytesIO:
    """Create a demo document that should pass '立项报告' checks."""
    content = """
# XX项目实施立项报告

## 封面
项目名称：XX项目
编制单位：技术部
日期：2026年5月

## 项目背景
随着公司业务发展，现有系统已无法满足需求。

## 投资估算
项目总预算：500万元
资金来源：公司自有资金
分项费用：
- 硬件采购：200万元
- 软件开发：200万元
- 实施服务：100万元

## 技术方案
采用微服务架构，基于Spring Cloud技术栈。
方案A：自建方案，成本可控但开发周期长
方案B：采购商业软件，快速上线但定制性差
对比结论：推荐方案A，长期成本更优

## 风险分析
主要风险包括技术风险、进度风险、人员风险，均已制定应对措施。
"""
    return make_docx(content)


def make_noncompliant_report_docx() -> BytesIO:
    """Create a demo document that should FAIL '立项报告' checks."""
    content = """
# 会议纪要

参会人员：张三、李四、王五
会议时间：2026年5月25日

## 讨论内容
1. 讨论了系统升级的可能性
2. 初步分析了现有问题
3. 决定进一步调研

## 待办事项
- 完成调研报告
- 确认预算
"""
    return make_docx(content)


def make_large_docx(word_count: int = 50000) -> BytesIO:
    """Create a large .docx with approximately word_count words."""
    content = "# 大文档测试\n\n" + "测试段落。 " * word_count
    return make_docx(content)


def make_encrypted_docx() -> BytesIO:
    """Create a password-protected .docx (not really encrypted, just flagged).

    Note: python-docx cannot create encrypted docs.
    We return a regular docx for cases where encryption check is server-side.
    """
    return make_docx("This is a test")
