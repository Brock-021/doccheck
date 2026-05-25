"""
DocCheck 报告导出服务

支持导出格式：PDF (reportlab) / Word (python-docx)
"""

import os
import io
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _conclusion_label(conclusion: str | None) -> str:
    labels = {
        "pass": "✅ 通过",
        "conditional_pass": "⚠️ 有条件通过",
        "fail": "❌ 不通过",
    }
    return labels.get(conclusion, "⏳ 待审核") if conclusion else "⏳ 待审核"


def _stage_label(stage: str) -> str:
    return {"initial": "初检", "final": "终检", "all": "通用"}.get(stage, stage)


def export_docx(
    report_data: dict,
    output_path: str | Path,
):
    """
    导出 Word (.docx) 格式检查报告。

    Args:
        report_data: API /api/reports/{id} 返回的数据
        output_path: 输出文件路径
    """
    doc = DocxDocument()

    # ── 标题 ──
    title = doc.add_heading("DocCheck 文档检查报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 文档信息 ──
    doc.add_heading("文档信息", level=1)
    doc_info = report_data.get("document", {})
    task_info = report_data.get("check_task", {})
    report_meta = report_data

    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = "Light Shading Accent 1"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows_data = [
        ("文件名", doc_info.get("filename", "-")),
        ("文档类型", doc_info.get("doc_type_name", "-")),
        ("上传人", doc_info.get("uploader", "-")),
        ("上传时间", doc_info.get("upload_time", "-")),
        ("检查阶段", _stage_label(task_info.get("stage", ""))),
        ("审核结论", _conclusion_label(report_meta.get("conclusion"))),
        ("结论说明", report_meta.get("conclusion_remark", "") or "无"),
    ]
    for i, (key, val) in enumerate(rows_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = str(val)

    doc.add_paragraph("")

    # ── 统计摘要 ──
    doc.add_heading("统计摘要", level=1)
    results = report_data.get("results", [])
    total = len(results)
    passed = sum(1 for r in results if r.get("compliant") == "true")
    failed = sum(1 for r in results if r.get("compliant") != "true")
    confirmed = sum(1 for r in results if r.get("review_status") == "confirmed")

    stats_table = doc.add_table(rows=2, cols=5)
    stats_table.style = "Light Shading Accent 1"
    headers = ["检查项", "通过", "问题", "已确认", "通过率"]
    stats_values = [str(total), str(passed), str(failed), str(confirmed),
                    f"{round(passed / total * 100, 1) if total > 0 else 0}%"]

    for i, h in enumerate(headers):
        stats_table.rows[0].cells[i].text = h
        stats_table.rows[1].cells[i].text = stats_values[i]

    doc.add_paragraph("")

    # ── 检查明细 ──
    doc.add_heading("检查明细", level=1)

    for r in results:
        compliant = r.get("compliant") == "true"
        severity = r.get("rule_severity", "suggest")
        severity_label = "必须改" if severity == "must_fix" else "建议改"
        status_icon = "✅ 通过" if compliant else "❌ 不通过"

        p = doc.add_paragraph()
        run = p.add_run(f"■ {r.get('rule_name', '未知规则')}  [{status_icon}]  ({severity_label})")
        run.bold = True

        if not compliant:
            if r.get("issue"):
                doc.add_paragraph(f"  问题: {r['issue']}")
            if r.get("location"):
                doc.add_paragraph(f"  位置: {r['location']}")
            if r.get("suggestion"):
                doc.add_paragraph(f"  建议: {r['suggestion']}")
            if r.get("review_status") == "confirmed":
                doc.add_paragraph(f"  审核: ✅ 已确认")
            elif r.get("review_status") == "rejected":
                doc.add_paragraph(f"  审核: ✏️ 已驳回" + (f" (备注: {r['review_remark']})" if r.get('review_remark') else ""))
            elif r.get("review_status") == "ignored":
                doc.add_paragraph(f"  审核: ⏭️ 已忽略" + (f" (备注: {r['review_remark']})" if r.get('review_remark') else ""))
            else:
                doc.add_paragraph(f"  审核: ⏳ 待审核")
        doc.add_paragraph("")

    # ── 页脚 ──
    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(f"DocCheck · 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(str(output_path))


def export_pdf(
    report_data: dict,
    output_path: str | Path,
):
    """
    导出 PDF 格式检查报告。

    使用 reportlab 生成，无需额外系统依赖。

    Args:
        report_data: API /api/reports/{id} 返回的数据
        output_path: 输出文件路径
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm, cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable,
    )

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    story = []

    # ── Title ──
    title_style = ParagraphStyle(
        "DocCheckTitle",
        parent=styles["Title"],
        fontSize=20,
        spaceAfter=6 * mm,
        textColor=HexColor("#1e293b"),
    )
    story.append(Paragraph("DocCheck 文档检查报告", title_style))
    story.append(Spacer(1, 3 * mm))

    # ── Document Info Table ──
    doc_info = report_data.get("document", {})
    task_info = report_data.get("check_task", {})
    report_meta = report_data

    info_data = [
        ["文件名", doc_info.get("filename", "-")],
        ["文档类型", doc_info.get("doc_type_name", "-")],
        ["上传人", doc_info.get("uploader", "-")],
        ["检查阶段", _stage_label(task_info.get("stage", ""))],
        ["审核结论", _conclusion_label(report_meta.get("conclusion"))],
        ["结论说明", report_meta.get("conclusion_remark", "") or "无"],
    ]

    info_table = Table(info_data, colWidths=[80 * mm, 80 * mm])
    info_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BACKGROUND", (0, 0), (0, -1), HexColor("#f1f5f9")),
        ("ALIGN", (0, 0), (0, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 6 * mm))

    # ── Summary ──
    story.append(Paragraph("统计摘要", styles["Heading2"]))
    results = report_data.get("results", [])
    total = len(results)
    passed = sum(1 for r in results if r.get("compliant") == "true")
    failed = sum(1 for r in results if r.get("compliant") != "true")
    confirmed = sum(1 for r in results if r.get("review_status") == "confirmed")
    pass_rate = f"{round(passed / total * 100, 1) if total > 0 else 0}%"

    stats_data = [
        ["检查项", "通过", "问题", "已确认", "通过率"],
        [str(total), str(passed), str(failed), str(confirmed), pass_rate],
    ]
    stats_table = Table(stats_data, colWidths=[40 * mm, 30 * mm, 30 * mm, 30 * mm, 30 * mm])
    stats_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#3b82f6")),
        ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(stats_table)
    story.append(Spacer(1, 6 * mm))

    # ── Check Details ──
    story.append(Paragraph("检查明细", styles["Heading2"]))

    for r in results:
        compliant = r.get("compliant") == "true"
        severity = r.get("rule_severity", "suggest")
        severity_label = "必须改" if severity == "must_fix" else "建议改"
        status_icon = "✅ 通过" if compliant else "❌ 不通过"

        # Rule name
        rule_text = f"<b>■ {r.get('rule_name', '未知规则')}  [{status_icon}]  ({severity_label})</b>"
        story.append(Paragraph(rule_text, styles["Normal"]))

        if not compliant:
            if r.get("issue"):
                story.append(Paragraph(f"&nbsp;&nbsp;问题: {r['issue']}", styles["Normal"]))
            if r.get("location"):
                story.append(Paragraph(f"&nbsp;&nbsp;位置: {r['location']}", styles["Normal"]))
            if r.get("suggestion"):
                story.append(Paragraph(f"&nbsp;&nbsp;建议: {r['suggestion']}", styles["Normal"]))

            review_status = r.get("review_status", "pending")
            if review_status == "confirmed":
                story.append(Paragraph("&nbsp;&nbsp;审核: ✅ 已确认", styles["Normal"]))
            elif review_status == "rejected":
                remark = r.get("review_remark", "")
                story.append(Paragraph(f"&nbsp;&nbsp;审核: ✏️ 已驳回{f' (备注: {remark})' if remark else ''}", styles["Normal"]))
            elif review_status == "ignored":
                remark = r.get("review_remark", "")
                story.append(Paragraph(f"&nbsp;&nbsp;审核: ⏭️ 已忽略{f' (备注: {remark})' if remark else ''}", styles["Normal"]))
            else:
                story.append(Paragraph("&nbsp;&nbsp;审核: ⏳ 待审核", styles["Normal"]))

        story.append(Spacer(1, 2 * mm))
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#e2e8f0")))
        story.append(Spacer(1, 2 * mm))

    # ── Footer ──
    story.append(Spacer(1, 10 * mm))
    footer_text = f"DocCheck · 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    story.append(Paragraph(footer_text, ParagraphStyle(
        "Footer",
        parent=styles["Normal"],
        fontSize=8,
        textColor=HexColor("#94a3b8"),
        alignment=1,
    )))

    doc.build(story)


def get_report_data_for_export(
    report_id: int,
    api_base_url: str,
    session_cookie: str,
) -> dict:
    """
    从 API 获取报告数据用于导出。

    Args:
        report_id: 报告 ID
        api_base_url: API 基础 URL，如 http://localhost:8001
        session_cookie: 用户的 session cookie 值

    Returns:
        report_data dict（与 /api/reports/{id} 返回格式一致）
    """
    import httpx
    resp = httpx.get(
        f"{api_base_url.rstrip('/')}/api/reports/{report_id}",
        cookies={"session": session_cookie},
        timeout=10,
    )
    resp.raise_for_status()
    return resp.json()
