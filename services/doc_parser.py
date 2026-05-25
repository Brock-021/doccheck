"""
DocCheck Word 文档解析服务
"""

from docx import Document
from pathlib import Path


def parse_docx(file_path: str | Path) -> dict:
    """解析 .docx 文件，返回文本内容和章节结构。

    Returns:
        {
            "full_text": "...",
            "sections": [
                {"level": 1, "title": "...", "text": "..."},
                ...
            ],
            "paragraph_count": N,
            "char_count": N,
        }
    """
    doc = Document(str(file_path))

    sections = []
    current_section = None
    full_text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        full_text_parts.append(text)

        if para.style.name.startswith("Heading"):
            if current_section:
                sections.append(current_section)
            # Determine level
            level = 1
            if "2" in para.style.name:
                level = 2
            elif "3" in para.style.name:
                level = 3
            current_section = {"level": level, "title": text, "text": ""}
        else:
            if current_section:
                current_section["text"] += text + "\n"

    # Last section
    if current_section:
        sections.append(current_section)

    full_text = "\n".join(full_text_parts)

    return {
        "full_text": full_text,
        "sections": sections,
        "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
        "char_count": len(full_text),
    }
